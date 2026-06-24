"""
file_manager.py
===============
Module quan ly file:
- Luu chu ky ra .sig (JSON – dinh dang chinh, chuan cryptography)
- Luu chu ky ra .txt
- Luu chu ky ra .docx
- Luu chu ky ra .csv (DataFrame format)
- Doc file chu ky
- [PDF] Ho tro .pdf (comment – bo comment de bat)
"""

import os
import csv
import json
from datetime import datetime
from pathlib import Path

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# ==============================================================================
#  LUU / DOC FILE .TXT (mac dinh)
# ==============================================================================

def save_as_txt(filepath: str, data: dict) -> None:
    """
    Luu thong tin chu ky ra file .txt.

    data dict chua cac key:
        message, hash_algorithm, hash_hex, signature_hex,
        public_key_e, public_key_n, private_key_d (tuy chon)
    """
    lines = [
        "=" * 60,
        "  UNG DUNG CHU KY SO RSA",
        "  An toan va Bao mat Thong tin",
        "=" * 60,
        f"  Thoi gian: {_timestamp()}",
        "=" * 60,
        "",
        "[THONG DIEP GOC]",
        data.get("message", ""),
        "",
        "[THUAT TOAN HASH]",
        data.get("hash_algorithm", "SHA-256"),
        "",
        "[GIA TRI HASH (H(m))]",
        data.get("hash_hex", ""),
        "",
        "[CHU KY SO (S = H^d mod n)]",
        data.get("signature_hex", ""),
        "",
        "[KHOA CONG KHAI]",
        f"  e = {data.get('public_key_e', '')}",
        f"  n = {data.get('public_key_n', '')}",
        "",
    ]

    if data.get("include_private", False):
        lines += [
            "[KHOA BI MAT]",
            f"  d = {data.get('private_key_d', '')}",
            f"  n = {data.get('public_key_n', '')}",
            "",
        ]

    lines += [
        "=" * 60,
        "  [Ket thuc file chu ky]",
        "=" * 60,
    ]

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ==============================================================================
#  LUU / DOC FILE .DOCX
# ==============================================================================

def save_as_docx(filepath: str, data: dict) -> None:
    """Luu thong tin chu ky ra file .docx."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("Thu vien python-docx chua duoc cai dat.")

    doc = Document()

    title = doc.add_heading("UNG DUNG CHU KY SO RSA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("An toan va Bao mat Thong tin")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Thoi gian tao: {_timestamp()}")
    doc.add_paragraph("-" * 50)

    def add_section(label: str, content: str):
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(content or "(trong)")
        doc.add_paragraph("")

    add_section("THONG DIEP GOC", data.get("message", ""))
    add_section("THUAT TOAN HASH", data.get("hash_algorithm", "SHA-256"))
    add_section("GIA TRI HASH  H(m)", data.get("hash_hex", ""))
    add_section("CHU KY SO  S = H^d mod n", data.get("signature_hex", ""))

    doc.add_paragraph("").add_run("KHOA CONG KHAI").bold = True
    doc.add_paragraph(f"  e = {data.get('public_key_e', '')}")
    doc.add_paragraph(f"  n = {data.get('public_key_n', '')}")

    if data.get("include_private", False):
        doc.add_paragraph("").add_run("KHOA BI MAT").bold = True
        doc.add_paragraph(f"  d = {data.get('private_key_d', '')}")
        doc.add_paragraph(f"  n = {data.get('public_key_n', '')}")

    doc.save(filepath)


# ==============================================================================
#  [PDF] HO TRO LUU FILE .PDF  –  yeu cau: pip install reportlab>=4.0.0
#
#  Huong dan kich hoat:
#    1. Cai dat thu vien:  pip install reportlab>=4.0.0
#    2. Bo comment toan bo khoi duoi (tu "PDF_SAVE_START" den "PDF_SAVE_END")
#    3. Bo comment 2 dong "elif ext == '.pdf'" ben duoi trong save_signature_file()
# ==============================================================================

# --- PDF_SAVE_START ---
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def save_as_pdf(filepath: str, data: dict) -> None:
    """
    Luu thong tin chu ky ra file .pdf su dung reportlab.

    Yeu cau: pip install reportlab>=4.0.0
    """
    if not PDF_AVAILABLE:
        raise RuntimeError(
            "Thu vien reportlab chua duoc cai dat.\n"
            "Chay: pip install reportlab>=4.0.0"
        )

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(
        "rsa_title",
        parent=styles["Title"],
        fontSize=16,
        textColor=colors.HexColor("#1a3a6b"),
        spaceAfter=4,
        alignment=1,
    )
    style_sub = ParagraphStyle(
        "rsa_sub",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#444444"),
        alignment=1,
        spaceAfter=6,
    )
    style_heading = ParagraphStyle(
        "rsa_heading",
        parent=styles["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1a3a6b"),
        spaceBefore=10,
        spaceAfter=2,
    )
    style_body = ParagraphStyle(
        "rsa_body",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Courier",
        textColor=colors.HexColor("#333333"),
        spaceAfter=4,
        wordWrap="CJK",
    )

    story = []
    story.append(Paragraph("UNG DUNG CHU KY SO RSA", style_title))
    story.append(Paragraph("An toan va Bao mat Thong tin", style_sub))
    story.append(Paragraph(f"Thoi gian tao: {_timestamp()}", style_sub))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 6))

    def add_pdf_section(label: str, value: str):
        story.append(Paragraph(label, style_heading))
        chunk_size = 80
        val = value or "(trong)"
        chunks = [val[i:i+chunk_size] for i in range(0, len(val), chunk_size)]
        for chunk in chunks:
            story.append(Paragraph(chunk, style_body))
        story.append(Spacer(1, 4))

    add_pdf_section("THONG DIEP GOC", data.get("message", ""))
    add_pdf_section("THUAT TOAN HASH", data.get("hash_algorithm", "SHA-256"))
    add_pdf_section("GIA TRI HASH  H(m)", data.get("hash_hex", ""))
    add_pdf_section("CHU KY SO  S = H^d mod n", data.get("signature_hex", ""))
    add_pdf_section("KHOA CONG KHAI  e", str(data.get("public_key_e", "")))
    add_pdf_section("KHOA CONG KHAI  n", str(data.get("public_key_n", "")))

    if data.get("include_private", False):
        add_pdf_section("KHOA BI MAT  d", str(data.get("private_key_d", "")))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Paragraph("[Ket thuc file chu ky]", style_sub))
    doc.build(story)
# --- PDF_SAVE_END ---


# --- PDF_READ_START ---
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False


def read_pdf(filepath: str) -> str:
    """
    Doc noi dung van ban tu file .pdf va tra ve chuoi.

    Yeu cau: pip install pypdf>=3.0.0
    """
    if not PYPDF_AVAILABLE:
        raise RuntimeError(
            "Thu vien pypdf chua duoc cai dat.\n"
            "Chay: pip install pypdf>=3.0.0"
        )
    reader = PdfReader(filepath)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text)
    return "\n".join(pages_text)
# --- PDF_READ_END ---


# ==============================================================================
#  LUU / DOC FILE .SIG  (JSON – dinh dang chuan cho chu ky so)
# ==============================================================================

def save_as_sig(filepath: str, data: dict) -> None:
    """
    Luu thong tin chu ky ra file .sig (dinh dang JSON).

    Day la dinh dang chuan trong linh vuc cryptography (tuong tu OpenSSL, GPG).
    Chua day du metadata: thong diep, hash, chu ky, khoa cong khai, thoi gian.

    data dict chua cac key:
        message, hash_algorithm, hash_hex, signature_hex,
        public_key_e, public_key_n, private_key_d (tuy chon)
    """
    payload = {
        "format":         "RSA-SIG-v1",
        "timestamp":      _timestamp(),
        "message":        data.get("message", ""),
        "hash_algorithm": data.get("hash_algorithm", "SHA-256"),
        "hash_hex":       data.get("hash_hex", ""),
        "signature_hex":  data.get("signature_hex", ""),
        "public_key": {
            "e": str(data.get("public_key_e", "")),
            "n": str(data.get("public_key_n", "")),
        },
    }

    if data.get("include_private", False):
        payload["private_key"] = {
            "d": str(data.get("private_key_d", "")),
            "n": str(data.get("public_key_n", "")),
        }

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def read_sig(filepath: str) -> dict:
    """
    Doc file .sig va tra ve dict chua toan bo thong tin chu ky.

    Returns:
        dict voi cac key: format, timestamp, message, hash_algorithm,
                          hash_hex, signature_hex, public_key, [private_key]
    """
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


# ==============================================================================
#  LUU / DOC FILE .CSV  (DataFrame format)
# ==============================================================================

# Danh sach cot mac dinh cho CSV chu ky
CSV_COLUMNS = [
    "timestamp",
    "message",
    "hash_algorithm",
    "hash_hex",
    "signature_hex",
    "public_key_e",
    "public_key_n",
    "private_key_d",
]


def save_as_csv(filepath: str, data: dict) -> None:
    """
    Luu thong tin chu ky ra file .csv theo dinh dang DataFrame.

    Moi lan ky se ghi them 1 dong (append) vao file CSV.
    Neu file chua ton tai, tao moi voi header.

    data dict chua cac key:
        message, hash_algorithm, hash_hex, signature_hex,
        public_key_e, public_key_n, private_key_d (tuy chon)
    """
    file_exists = Path(filepath).exists()

    row = {
        "timestamp":      _timestamp(),
        "message":        data.get("message", ""),
        "hash_algorithm": data.get("hash_algorithm", ""),
        "hash_hex":       data.get("hash_hex", ""),
        "signature_hex":  data.get("signature_hex", ""),
        "public_key_e":   str(data.get("public_key_e", "")),
        "public_key_n":   str(data.get("public_key_n", "")),
        "private_key_d":  str(data.get("private_key_d", "")) if data.get("include_private") else "",
    }

    with open(filepath, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=CSV_COLUMNS)
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)


def read_signature_csv(filepath: str):
    """
    Doc file CSV chu ky va tra ve DataFrame (neu co pandas)
    hoac list[dict] (neu khong co pandas).

    Returns:
        pandas.DataFrame  khi PANDAS_AVAILABLE == True
        list[dict]         khi pandas chua duoc cai dat
    """
    if PANDAS_AVAILABLE:
        df = pd.read_csv(filepath, dtype=str).fillna("")
        return df
    else:
        rows = []
        with open(filepath, "r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(dict(row))
        return rows


# ==============================================================================
#  HAM CONG KHAI – tu dong chon dinh dang theo extension
# ==============================================================================

def save_signature_file(filepath: str, data: dict) -> None:
    """Luu file chu ky (tu dong chon dinh dang theo extension)."""
    ext = Path(filepath).suffix.lower()
    if ext == ".sig":
        save_as_sig(filepath, data)
    elif ext == ".docx":
        save_as_docx(filepath, data)
    elif ext == ".csv":
        save_as_csv(filepath, data)
    # Bo comment 2 dong duoi de ho tro luu ra .pdf (can reportlab):
    elif ext == ".pdf":
        save_as_pdf(filepath, data)
    else:
        save_as_txt(filepath, data)


def read_file(filepath: str) -> str:
    """Doc noi dung file (txt, docx, sig, hoac pdf) va tra ve chuoi."""
    ext = Path(filepath).suffix.lower()
    if ext == ".sig":
        # Tra ve signature_hex tu file .sig (JSON)
        sig_data = read_sig(filepath)
        return sig_data.get("signature_hex", "")
    elif ext == ".docx":
        if not DOCX_AVAILABLE:
            raise RuntimeError("Thu vien python-docx chua duoc cai dat.")
        doc = Document(filepath)
        lines = [para.text for para in doc.paragraphs]
        return "\n".join(lines)
    elif ext == ".csv":
        # Tra ve noi dung signature_hex cua dong cuoi cung (de tuong thich nguoc)
        rows = read_signature_csv(filepath)
        if PANDAS_AVAILABLE:
            import pandas as pd
            if isinstance(rows, pd.DataFrame) and not rows.empty:
                return rows.iloc[-1].get("signature_hex", "")
        else:
            if rows:
                return rows[-1].get("signature_hex", "")
        return ""
    # Bo comment 2 dong duoi de ho tro doc file .pdf (can pypdf):
    elif ext == ".pdf":
        return read_pdf(filepath)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
